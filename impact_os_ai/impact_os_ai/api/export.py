import frappe
from frappe import _
from .auth import get_current_user_from_token
import json
import io
import base64
from datetime import datetime


def _get_project_and_sections(slug: str, user_email: str):
    """Fetch project doc and parsed sections dict, checking authorization."""
    if not frappe.db.exists("IOS Project", {"slug": slug}):
        frappe.throw(_("Project not found"), frappe.DoesNotExistError)

    doc = frappe.get_doc("IOS Project", {"slug": slug})

    if doc.owner != user_email and "System Manager" not in frappe.get_roles(user_email):
        frappe.throw(_("Not authorized to export this project"), frappe.PermissionError)

    sections = {}
    if doc.generated_sections:
        try:
            sections = json.loads(doc.generated_sections)
        except Exception:
            sections = {}

    if not sections:
        frappe.throw(
            _("No generated content found. Please generate MAP sections before exporting."),
            frappe.ValidationError,
        )

    return doc, sections


@frappe.whitelist(allow_guest=True)
def export_docx(slug: str):
    """
    Export project as a Word document (.docx).
    GET /api/method/impact_os_ai.impact_os_ai.api.export.export_docx?slug=xxx
    """
    user_email = get_current_user_from_token()
    doc, sections = _get_project_and_sections(slug, user_email)

    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        frappe.throw(_("python-docx is not installed. Please install it."), frappe.ValidationError)

    document = DocxDocument()

    # Title page
    title_para = document.add_heading(doc.project_title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_para = document.add_paragraph(f"{doc.organization} | {doc.sector}")
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = document.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # Project overview
    document.add_heading("Project Overview", level=1)
    overview_table = document.add_table(rows=6, cols=2)
    overview_table.style = "Table Grid"
    rows_data = [
        ("Organization", doc.organization),
        ("Sector", doc.sector),
        ("Country/Region", doc.country or "N/A"),
        ("Budget (USD)", f"${doc.budget_usd:,.0f}" if doc.budget_usd else "N/A"),
        ("Duration", f"{doc.duration_months or 12} months"),
        ("Target Beneficiaries", str(doc.target_beneficiaries or "N/A")),
    ]
    for i, (label, value) in enumerate(rows_data):
        overview_table.rows[i].cells[0].text = label
        overview_table.rows[i].cells[1].text = value

    if doc.description:
        document.add_heading("Project Description", level=2)
        document.add_paragraph(doc.description)

    document.add_page_break()

    # Sections ordered by MAP framework
    section_order = [
        "mission_vision", "theory_of_change", "logframe", "stakeholder_map",
        "indicators", "data_collection", "budget_narrative", "risk_matrix",
        "evaluation_plan", "sustainability",
    ]

    for section_key in section_order:
        if section_key not in sections:
            continue
        section_data = sections[section_key]
        document.add_heading(section_data.get("label", section_key.replace("_", " ").title()), level=1)

        content = section_data.get("content", "")
        for paragraph in content.split("\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            if paragraph.startswith("# "):
                document.add_heading(paragraph[2:], level=2)
            elif paragraph.startswith("## "):
                document.add_heading(paragraph[3:], level=3)
            elif paragraph.startswith("### "):
                document.add_heading(paragraph[4:], level=4)
            elif paragraph.startswith("- ") or paragraph.startswith("* "):
                p = document.add_paragraph(paragraph[2:], style="List Bullet")
            elif paragraph.startswith("**") and paragraph.endswith("**"):
                p = document.add_paragraph()
                run = p.add_run(paragraph[2:-2])
                run.bold = True
            else:
                document.add_paragraph(paragraph)

        document.add_page_break()

    # Save to buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    content_b64 = base64.b64encode(buffer.read()).decode("utf-8")

    _log_export(user_email, slug, "docx")

    filename = f"{doc.slug}-mis-blueprint.docx"
    return {
        "filename": filename,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "data": content_b64,
    }


@frappe.whitelist(allow_guest=True)
def export_excel(slug: str):
    """
    Export project as an Excel workbook (.xlsx).
    GET /api/method/impact_os_ai.impact_os_ai.api.export.export_excel?slug=xxx
    """
    user_email = get_current_user_from_token()
    doc, sections = _get_project_and_sections(slug, user_email)

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        frappe.throw(_("openpyxl is not installed. Please install it."), frappe.ValidationError)

    wb = openpyxl.Workbook()

    # Header styles
    header_font = Font(name="Calibri", bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Overview sheet
    ws_overview = wb.active
    ws_overview.title = "Project Overview"
    ws_overview.column_dimensions["A"].width = 30
    ws_overview.column_dimensions["B"].width = 60

    overview_data = [
        ("Project Title", doc.project_title),
        ("Organization", doc.organization),
        ("Sector", doc.sector),
        ("Country/Region", doc.country or "N/A"),
        ("Budget (USD)", f"${doc.budget_usd:,.0f}" if doc.budget_usd else "N/A"),
        ("Duration", f"{doc.duration_months or 12} months"),
        ("Target Beneficiaries", str(doc.target_beneficiaries or "N/A")),
        ("SDG Goals", doc.sdg_goals or "N/A"),
        ("Generated Date", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Description", doc.description or "N/A"),
    ]

    for row_idx, (label, value) in enumerate(overview_data, start=1):
        cell_a = ws_overview.cell(row=row_idx, column=1, value=label)
        cell_a.font = Font(bold=True)
        cell_b = ws_overview.cell(row=row_idx, column=2, value=value)
        cell_b.alignment = Alignment(wrap_text=True)

    # One sheet per section
    section_order = [
        "mission_vision", "theory_of_change", "logframe", "stakeholder_map",
        "indicators", "data_collection", "budget_narrative", "risk_matrix",
        "evaluation_plan", "sustainability",
    ]

    for section_key in section_order:
        if section_key not in sections:
            continue
        section_data = sections[section_key]
        sheet_name = section_data.get("label", section_key)[:31]  # Excel tab limit
        ws = wb.create_sheet(title=sheet_name)

        ws.column_dimensions["A"].width = 120
        ws.merge_cells("A1:A1")
        header_cell = ws["A1"]
        header_cell.value = section_data.get("label", section_key)
        header_cell.font = header_font
        header_cell.fill = header_fill
        header_cell.alignment = header_alignment
        ws.row_dimensions[1].height = 30

        content = section_data.get("content", "")
        row = 2
        for line in content.split("\n"):
            cell = ws.cell(row=row, column=1, value=line)
            cell.alignment = Alignment(wrap_text=True)
            ws.row_dimensions[row].height = 15
            row += 1

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    content_b64 = base64.b64encode(buffer.read()).decode("utf-8")

    _log_export(user_email, slug, "xlsx")

    filename = f"{doc.slug}-mis-blueprint.xlsx"
    return {
        "filename": filename,
        "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "data": content_b64,
    }


@frappe.whitelist(allow_guest=True)
def export_pdf(slug: str):
    """
    Export project as PDF (HTML-based, rendered by Frappe's PDF engine).
    GET /api/method/impact_os_ai.impact_os_ai.api.export.export_pdf?slug=xxx
    """
    user_email = get_current_user_from_token()
    doc, sections = _get_project_and_sections(slug, user_email)

    html = _build_html(doc, sections)

    try:
        from frappe.utils.pdf import get_pdf
        pdf_data = get_pdf(html)
        content_b64 = base64.b64encode(pdf_data).decode("utf-8")
    except Exception as e:
        frappe.log_error(f"PDF generation error: {str(e)}", "ImpactOS Export PDF")
        # Fall back to returning HTML as base64
        content_b64 = base64.b64encode(html.encode("utf-8")).decode("utf-8")
        _log_export(user_email, slug, "html")
        return {
            "filename": f"{doc.slug}-mis-blueprint.html",
            "content_type": "text/html",
            "data": content_b64,
        }

    _log_export(user_email, slug, "pdf")
    return {
        "filename": f"{doc.slug}-mis-blueprint.pdf",
        "content_type": "application/pdf",
        "data": content_b64,
    }


@frappe.whitelist(allow_guest=True)
def get_export_history(slug: str = ""):
    """
    Get export log history.
    GET /api/method/impact_os_ai.impact_os_ai.api.export.get_export_history
    """
    user_email = get_current_user_from_token()

    filters = {"user": user_email}
    if slug:
        filters["project_slug"] = slug

    logs = frappe.get_all(
        "IOS Export Log",
        filters=filters,
        fields=["name", "project_slug", "export_format", "creation"],
        order_by="creation desc",
        limit_page_length=50,
    )
    return {"exports": logs}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _log_export(user_email: str, slug: str, fmt: str):
    try:
        log = frappe.get_doc({
            "doctype": "IOS Export Log",
            "user": user_email,
            "project_slug": slug,
            "export_format": fmt,
        })
        log.insert(ignore_permissions=True)
        frappe.db.commit()
    except Exception as e:
        frappe.log_error(f"Failed to log export: {str(e)}", "ImpactOS Export Log")


def _build_html(doc, sections: dict) -> str:
    section_order = [
        "mission_vision", "theory_of_change", "logframe", "stakeholder_map",
        "indicators", "data_collection", "budget_narrative", "risk_matrix",
        "evaluation_plan", "sustainability",
    ]

    sections_html = ""
    for key in section_order:
        if key not in sections:
            continue
        sd = sections[key]
        content_html = sd.get("content", "").replace("\n", "<br>")
        sections_html += f"""
        <div class="section">
            <h2>{sd.get('label', key)}</h2>
            <div class="content">{content_html}</div>
        </div>
        """

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{doc.project_title} — MIS Blueprint</title>
<style>
  body {{ font-family: 'Calibri', Arial, sans-serif; margin: 40px; color: #222; }}
  h1 {{ color: #1F4E79; text-align: center; }}
  h2 {{ color: #2E74B5; border-bottom: 2px solid #2E74B5; padding-bottom: 6px; margin-top: 40px; }}
  .meta {{ text-align: center; color: #666; margin-bottom: 30px; }}
  .section {{ page-break-inside: avoid; margin-bottom: 40px; }}
  .content {{ line-height: 1.7; }}
  table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
  td, th {{ border: 1px solid #ccc; padding: 8px 12px; }}
  th {{ background: #1F4E79; color: white; }}
</style>
</head>
<body>
<h1>{doc.project_title}</h1>
<p class="meta">{doc.organization} | {doc.sector} | Generated: {datetime.now().strftime('%B %d, %Y')}</p>

<div class="section">
  <h2>Project Overview</h2>
  <table>
    <tr><th>Field</th><th>Details</th></tr>
    <tr><td>Organization</td><td>{doc.organization}</td></tr>
    <tr><td>Sector</td><td>{doc.sector}</td></tr>
    <tr><td>Country/Region</td><td>{doc.country or 'N/A'}</td></tr>
    <tr><td>Budget (USD)</td><td>${doc.budget_usd:,.0f}</td></tr>
    <tr><td>Duration</td><td>{doc.duration_months or 12} months</td></tr>
    <tr><td>Target Beneficiaries</td><td>{doc.target_beneficiaries or 'N/A'}</td></tr>
    <tr><td>SDG Goals</td><td>{doc.sdg_goals or 'N/A'}</td></tr>
  </table>
  {f'<p>{doc.description}</p>' if doc.description else ''}
</div>

{sections_html}
</body>
</html>"""
